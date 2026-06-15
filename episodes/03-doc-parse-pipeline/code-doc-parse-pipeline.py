"""
7层文档解析 Pipeline — 尽调系统实战

来源：费曼学 AI 造物车间 #03（视频 24 分钟）
实测场景：私募尽调项目，每天 200+ 文档（PDF/扫描件/Excel/Word），要求解析后保留表格结构和原文行号。

7 层结构：
  L1 Schema 定义 → 统一 Document / Chunk 数据结构
  L2 Parser 路由  → 7路 parser（按文件类型 + 复杂度分）
  L3 PDF 4 档     → pdf-dots / pdfplumber / PyMuPDF / MinerU
  L4 OCR 4 降级   → Tesseract / PaddleOCR / Surya / VLM
  L5 Excel 多表   → openpyxl 多 sheet 解析
  L6 Word 统一    → python-docx 段落 + 表格
  L7 业务验收     → 表格行号 + 业务关键字校验

用法：
  from code_doc_parse_pipeline import parse_document
  doc = parse_document("path/to/file.pdf")
  for chunk in doc.chunks:
      print(f"[{chunk.page}] {chunk.text[:80]}")
"""

from __future__ import annotations

import logging
import re
import time
from abc import ABC, abstractmethod
from dataclasses import dataclass, field
from enum import Enum
from pathlib import Path
from typing import Any

logger = logging.getLogger(__name__)


# =============================================================================
# L1 — Schema 定义
# =============================================================================

class FileType(str, Enum):
    PDF_NATIVE = "pdf_native"          # 数字 PDF（可复制文字）
    PDF_SCAN = "pdf_scan"              # 扫描件 PDF（图片）
    EXCEL = "excel"                    # xlsx/xls
    WORD = "word"                      # docx
    TXT = "txt"
    IMAGE = "image"                    # 单图（PNG/JPG）
    UNKNOWN = "unknown"


class ParseStrategy(str, Enum):
    """L3 PDF 4 档"""
    PDF_DOTS = "pdf_dots"              # 极简：快速抽文本（<0.5s/页）
    PDFPLUMBER = "pdfplumber"          # 基础：保留表格位置（1-2s/页）
    PYMUPDF = "pymupdf"                # 中等：混合排版（2-3s/页）
    MINERU = "mineru"                  # 重型：扫描件 OCR + 旋转校正（10s/页）


@dataclass
class Chunk:
    """统一 chunk 数据结构"""
    text: str
    page: int = 0
    bbox: tuple[float, float, float, float] | None = None  # (x0, y0, x1, y1)
    table_id: str | None = None        # 表格 ID（如果有）
    line_no: int | None = None         # 原始行号（验收用）
    confidence: float = 1.0            # OCR 置信度（0-1）
    metadata: dict[str, Any] = field(default_factory=dict)


@dataclass
class Document:
    """统一 document 数据结构"""
    file_path: str
    file_type: FileType
    strategy: ParseStrategy | None = None
    chunks: list[Chunk] = field(default_factory=list)
    parse_time_ms: int = 0
    page_count: int = 0
    error: str | None = None
    metadata: dict[str, Any] = field(default_factory=dict)

    @property
    def text(self) -> str:
        """全文拼接（向后兼容）"""
        return "\n".join(c.text for c in self.chunks)

    @property
    def table_chunks(self) -> list[Chunk]:
        """仅返回表格 chunk"""
        return [c for c in self.chunks if c.table_id]

    def validate(self, keywords: list[str]) -> dict[str, Any]:
        """L7 业务验收：检查关键业务字段是否被正确解析"""
        full_text = self.text
        missing = [k for k in keywords if k not in full_text]
        return {
            "passed": len(missing) == 0,
            "missing_keywords": missing,
            "total_keywords": len(keywords),
            "found_keywords": len(keywords) - len(missing),
            "chunk_count": len(self.chunks),
            "table_count": len(self.table_chunks),
        }


# =============================================================================
# L2 — Parser 抽象基类
# =============================================================================

class BaseParser(ABC):
    """所有 parser 的统一接口"""

    file_types: list[FileType] = []
    priority: int = 100  # 数字越小越先尝试

    @abstractmethod
    def can_parse(self, path: Path) -> bool:
        """判断是否能处理该文件"""
        ...

    @abstractmethod
    def parse(self, path: Path) -> Document:
        """解析文件，返回 Document"""
        ...


# =============================================================================
# L3 — PDF 4 档策略
# =============================================================================

class PDFParser(BaseParser):
    """PDF 解析的基类，4 档策略共用"""
    file_types = [FileType.PDF_NATIVE, FileType.PDF_SCAN]
    priority = 50

    def can_parse(self, path: Path) -> bool:
        return path.suffix.lower() == ".pdf"

    def _detect_native_vs_scan(self, path: Path) -> FileType:
        """检测是数字 PDF 还是扫描件"""
        try:
            import pdfplumber
            with pdfplumber.open(path) as pdf:
                first_page = pdf.pages[0]
                # 抽前 5 页的平均字符数：>50 是数字 PDF
                total_chars = sum(len(p.extract_text() or "") for p in pdf.pages[:5])
                avg = total_chars / min(5, len(pdf.pages))
                return FileType.PDF_NATIVE if avg > 50 else FileType.PDF_SCAN
        except Exception:
            return FileType.PDF_NATIVE  # 默认按数字 PDF 处理

    def _is_complex_layout(self, path: Path) -> bool:
        """检测是否复杂排版（多栏 / 旋转 / 嵌入图片）"""
        try:
            import pdfplumber
            with pdfplumber.open(path) as pdf:
                first_page = pdf.pages[0]
                # 启发式：图片占比 > 30% 或 检测到旋转文字
                bbox_area = first_page.width * first_page.height
                image_area = sum(
                    (img.get("x1", 0) - img.get("x0", 0)) *
                    (img.get("bottom", 0) - img.get("top", 0))
                    for img in first_page.images
                )
                return image_area / bbox_area > 0.3
        except Exception:
            return False


class PDFDotsParser(PDFParser):
    """L3 档 1 — 极简：pypdf 抽纯文本，0.5s/页"""
    priority = 10
    strategy = ParseStrategy.PDF_DOTS

    def parse(self, path: Path) -> Document:
        start = time.time()
        chunks = []
        page_count = 0
        try:
            from pypdf import PdfReader
            reader = PdfReader(str(path))
            page_count = len(reader.pages)
            for i, page in enumerate(reader.pages):
                text = page.extract_text() or ""
                chunks.append(Chunk(text=text, page=i + 1, confidence=1.0))
        except Exception as e:
            return Document(
                file_path=str(path), file_type=FileType.PDF_NATIVE,
                error=f"pypdf failed: {e}", parse_time_ms=int((time.time() - start) * 1000),
            )
        return Document(
            file_path=str(path),
            file_type=FileType.PDF_NATIVE,
            strategy=self.strategy,
            chunks=chunks,
            page_count=page_count,
            parse_time_ms=int((time.time() - start) * 1000),
        )


class PDFPlumberParser(PDFParser):
    """L3 档 2 — 基础：pdfplumber 保留表格位置，1-2s/页"""
    priority = 20
    strategy = ParseStrategy.PDFPLUMBER

    def parse(self, path: Path) -> Document:
        start = time.time()
        chunks = []
        page_count = 0
        try:
            import pdfplumber
            with pdfplumber.open(path) as pdf:
                page_count = len(pdf.pages)
                for i, page in enumerate(pdf.pages):
                    text = page.extract_text() or ""
                    chunks.append(Chunk(text=text, page=i + 1, confidence=1.0))
                    # 提取表格
                    tables = page.extract_tables()
                    for t_idx, table in enumerate(tables):
                        table_id = f"p{i+1}_t{t_idx}"
                        # 表格 → markdown
                        md_lines = []
                        for row in table:
                            md_lines.append("| " + " | ".join(str(c or "") for c in row) + " |")
                        table_text = "\n".join(md_lines)
                        chunks.append(Chunk(
                            text=table_text,
                            page=i + 1,
                            table_id=table_id,
                            confidence=1.0,
                        ))
        except Exception as e:
            return Document(
                file_path=str(path), file_type=FileType.PDF_NATIVE,
                error=f"pdfplumber failed: {e}", parse_time_ms=int((time.time() - start) * 1000),
            )
        return Document(
            file_path=str(path),
            file_type=FileType.PDF_NATIVE,
            strategy=self.strategy,
            chunks=chunks,
            page_count=page_count,
            parse_time_ms=int((time.time() - start) * 1000),
        )


class PyMuPDFParser(PDFParser):
    """L3 档 3 — 中等：PyMuPDF 处理复杂排版 + 旋转，2-3s/页"""
    priority = 30
    strategy = ParseStrategy.PYMUPDF

    def parse(self, path: Path) -> Document:
        start = time.time()
        chunks = []
        page_count = 0
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(str(path))
            page_count = len(doc)
            for i, page in enumerate(doc):
                # 抽取文字块（保留位置）
                blocks = page.get_text("dict")["blocks"]
                for b_idx, block in enumerate(blocks):
                    if block.get("type") == 0:  # 文本块
                        text_lines = []
                        for line in block.get("lines", []):
                            for span in line.get("spans", []):
                                text_lines.append(span["text"])
                        if text_lines:
                            chunks.append(Chunk(
                                text=" ".join(text_lines),
                                page=i + 1,
                                bbox=tuple(block["bbox"]),
                                confidence=1.0,
                            ))
                # 提取表格（PyMuPDF 内置）
                try:
                    tabs = page.find_tables()
                    for t_idx, tab in enumerate(tabs):
                        df = tab.to_pandas()
                        table_id = f"p{i+1}_t{t_idx}"
                        chunks.append(Chunk(
                            text=df.to_markdown(index=False),
                            page=i + 1,
                            table_id=table_id,
                            confidence=1.0,
                        ))
                except Exception:
                    pass  # 没找到表格不算错误
            doc.close()
        except Exception as e:
            return Document(
                file_path=str(path), file_type=FileType.PDF_NATIVE,
                error=f"PyMuPDF failed: {e}", parse_time_ms=int((time.time() - start) * 1000),
            )
        return Document(
            file_path=str(path),
            file_type=FileType.PDF_NATIVE,
            strategy=self.strategy,
            chunks=chunks,
            page_count=page_count,
            parse_time_ms=int((time.time() - start) * 1000),
        )


class MinerUParser(PDFParser):
    """L3 档 4 — 重型：MinerU 处理扫描件 + 旋转 + 公式，10s/页"""
    priority = 40
    strategy = ParseStrategy.MINERU

    def parse(self, path: Path) -> Document:
        start = time.time()
        chunks = []
        page_count = 0
        try:
            # MinerU 提供高层 API（magic_pdf 库）
            from magic_pdf.pipe.UNIPipe import UNIPipe
            from magic_pdf.rw.DocsReader import DocsReader

            # 注意：MinerU 实际 API 复杂，这里给个骨架
            # 实际使用时参考 https://github.com/opendatalab/MinerU
            pipe = UNIPipe(
                DocsReader()(str(path)),
                os.environ.get("MINERU_MODEL_DIR", "/models/MinerU"),
            )
            pipe.pipe_classify()
            pipe.pipe_analyze()
            pipe.pipe_parse()

            page_count = len(pipe.doc.pages)
            for i, page in enumerate(pipe.doc.pages):
                # 抽取文本块（OCR 后的）
                for block in page.get_text_blocks():
                    chunks.append(Chunk(
                        text=block.text,
                        page=i + 1,
                        bbox=tuple(block.bbox),
                        confidence=block.conf,
                    ))
                # 抽取表格
                for t_idx, table in enumerate(page.get_tables()):
                    chunks.append(Chunk(
                        text=table.to_markdown(),
                        page=i + 1,
                        table_id=f"p{i+1}_t{t_idx}",
                        confidence=table.conf,
                    ))
        except ImportError:
            # MinerU 未安装 → fallback 到 4 级 OCR 降级（L4）
            logger.warning("MinerU 未安装，fallback 到 OCR 降级")
            return OCRPipeline().parse(path)
        except Exception as e:
            return Document(
                file_path=str(path), file_type=FileType.PDF_SCAN,
                error=f"MinerU failed: {e}", parse_time_ms=int((time.time() - start) * 1000),
            )
        return Document(
            file_path=str(path),
            file_type=FileType.PDF_SCAN,
            strategy=self.strategy,
            chunks=chunks,
            page_count=page_count,
            parse_time_ms=int((time.time() - start) * 1000),
        )


# =============================================================================
# L4 — OCR 4 降级（独立可用，也可被 L3 档 4 fallback）
# =============================================================================

class OCRLevel(int, Enum):
    L1_TESSERACT = 1   # 0 元/页，英文好
    L2_PADDLE = 2      # 0 元/页，中文好
    L3_SURYA = 3       # 0 元/页，旋转/复杂好
    L4_VLM = 4         # $0.01-0.03/页，复杂版式终极方案


class OCRPipeline:
    """4 级降级：每级失败自动升下一级"""

    def __init__(self, max_level: OCRLevel = OCRLevel.L4_VLM):
        self.max_level = max_level

    def parse(self, path: Path) -> Document:
        """主入口：按 OCRLevel 顺序降级"""
        start = time.time()
        chunks = []
        last_error = None
        for level in OCRLevel:
            if level > self.max_level:
                break
            try:
                if level == OCRLevel.L1_TESSERACT:
                    return self._tesseract(path)
                elif level == OCRLevel.L2_PADDLE:
                    return self._paddleocr(path)
                elif level == OCRLevel.L3_SURYA:
                    return self._surya(path)
                elif level == OCRLevel.L4_VLM:
                    return self._vlm(path)
            except Exception as e:
                last_error = f"L{level} failed: {e}"
                logger.warning(f"OCR L{level} 失败，降级到下一级: {e}")
                continue
        return Document(
            file_path=str(path),
            file_type=FileType.PDF_SCAN,
            error=f"All OCR levels failed. Last error: {last_error}",
            parse_time_ms=int((time.time() - start) * 1000),
        )

    def _tesseract(self, path: Path) -> Document:
        """L1 Tesseract：免费，英文好，中文一般"""
        # 实际使用：pytesseract + pdf2image
        # 这里给骨架
        raise NotImplementedError("L1 Tesseract: 集成 pytesseract + pdf2image")

    def _paddleocr(self, path: Path) -> Document:
        """L2 PaddleOCR：免费，中文好"""
        # 实际使用：paddleocr 库
        raise NotImplementedError("L2 PaddleOCR: 集成 paddleocr")

    def _surya(self, path: Path) -> Document:
        """L3 Surya：免费，旋转/复杂版式好"""
        # 实际使用：surya-ocr 库
        raise NotImplementedError("L3 Surya: 集成 surya-ocr")

    def _vlm(self, path: Path) -> Document:
        """L4 VLM：付费，复杂版式终极方案"""
        # 实际使用：GPT-4V / Claude / Qwen-VL
        raise NotImplementedError("L4 VLM: 集成 GPT-4V 或 Claude Vision")


# =============================================================================
# L5 — Excel 多表解析
# =============================================================================

class ExcelParser(BaseParser):
    """Excel 多 sheet 解析：保留 sheet 名 + 单元格坐标 + 公式"""
    file_types = [FileType.EXCEL]
    priority = 30

    def can_parse(self, path: Path) -> bool:
        return path.suffix.lower() in (".xlsx", ".xls")

    def parse(self, path: Path) -> Document:
        start = time.time()
        chunks = []
        try:
            import openpyxl
            wb = openpyxl.load_workbook(path, data_only=True)
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                rows = list(ws.iter_rows(values_only=False))
                if not rows:
                    continue

                # 标题行（第 1 行）
                headers = [c.value for c in rows[0]]

                # 数据行（第 2 行起）
                md_lines = ["| " + " | ".join(str(h or "") for h in headers) + " |"]
                md_lines.append("| " + " | ".join(["---"] * len(headers)) + " |")

                for row_idx, row in enumerate(rows[1:], start=2):
                    cells = []
                    for col_idx, cell in enumerate(row):
                        # 保留原始行号
                        cells.append(f"{cell.value or ''}")
                    md_lines.append("| " + " | ".join(cells) + " |")

                    # 重要单元格单独抽 chunk（带行号）
                    for col_idx, cell in enumerate(row):
                        if cell.value and self._is_important_cell(str(cell.value)):
                            chunks.append(Chunk(
                                text=f"[{sheet_name}!{cell.coordinate}] {cell.value}",
                                line_no=row_idx,
                                table_id=f"{sheet_name}",
                                confidence=1.0,
                            ))

                chunks.append(Chunk(
                    text="\n".join(md_lines),
                    table_id=sheet_name,
                    confidence=1.0,
                ))
        except Exception as e:
            return Document(
                file_path=str(path), file_type=FileType.EXCEL,
                error=f"openpyxl failed: {e}", parse_time_ms=int((time.time() - start) * 1000),
            )
        return Document(
            file_path=str(path),
            file_type=FileType.EXCEL,
            chunks=chunks,
            parse_time_ms=int((time.time() - start) * 1000),
        )

    def _is_important_cell(self, value: str) -> bool:
        """判断是否重要单元格：金额、日期、比例"""
        # 含数字 + 千分位 → 金额
        if re.search(r"\d{1,3}(,\d{3})+", value):
            return True
        # 含百分号 → 比例
        if "%" in value:
            return True
        # 含日期格式
        if re.search(r"\d{4}[-/]\d{1,2}[-/]\d{1,2}", value):
            return True
        return False


# =============================================================================
# L6 — Word 统一解析
# =============================================================================

class WordParser(BaseParser):
    """Word docx 解析：段落 + 表格 + 标题层级"""
    file_types = [FileType.WORD]
    priority = 30

    def can_parse(self, path: Path) -> bool:
        return path.suffix.lower() == ".docx"

    def parse(self, path: Path) -> Document:
        start = time.time()
        chunks = []
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(str(path))

            # 段落
            for i, para in enumerate(doc.paragraphs):
                if para.text.strip():
                    chunks.append(Chunk(
                        text=para.text,
                        line_no=i + 1,
                        confidence=1.0,
                    ))

            # 表格
            for t_idx, table in enumerate(doc.tables):
                md_lines = []
                for row in table.rows:
                    md_lines.append("| " + " | ".join(c.text for c in row.cells) + " |")
                chunks.append(Chunk(
                    text="\n".join(md_lines),
                    table_id=f"t{t_idx}",
                    confidence=1.0,
                ))
        except Exception as e:
            return Document(
                file_path=str(path), file_type=FileType.WORD,
                error=f"python-docx failed: {e}", parse_time_ms=int((time.time() - start) * 1000),
            )
        return Document(
            file_path=str(path),
            file_type=FileType.WORD,
            chunks=chunks,
            parse_time_ms=int((time.time() - start) * 1000),
        )


# =============================================================================
# L7 — 业务验收（实际只是 validate() 方法，前面 Schema 已有）
# =============================================================================

# 验收在 Document.validate() 实现，参考：
#   keywords = ["营业收入", "净利润", "ROE"]
#   result = doc.validate(keywords)
#   if not result["passed"]:
#       print(f"Missing: {result['missing_keywords']}")


# =============================================================================
# 路由 + 编排主入口
# =============================================================================

# Parser 注册表（按 priority 升序尝试）
PARSERS: list[BaseParser] = [
    PDFDotsParser(),
    PDFPlumberParser(),
    PyMuPDFParser(),
    MinerUParser(),
    ExcelParser(),
    WordParser(),
]


def detect_file_type(path: Path) -> FileType:
    """根据扩展名判断文件类型"""
    ext = path.suffix.lower()
    if ext == ".pdf":
        return FileType.PDF_NATIVE  # 默认按数字 PDF，后面 PDFParser 会再细分
    elif ext in (".xlsx", ".xls"):
        return FileType.EXCEL
    elif ext == ".docx":
        return FileType.WORD
    elif ext in (".txt", ".md"):
        return FileType.TXT
    elif ext in (".png", ".jpg", ".jpeg"):
        return FileType.IMAGE
    return FileType.UNKNOWN


def route_parser(path: Path, hint: ParseStrategy | None = None) -> BaseParser:
    """L2 路由：按 hint 或文件类型选 parser"""
    if hint == ParseStrategy.PDF_DOTS:
        return PDFDotsParser()
    elif hint == ParseStrategy.PDFPLUMBER:
        return PDFPlumberParser()
    elif hint == ParseStrategy.PYMUPDF:
        return PyMuPDFParser()
    elif hint == ParseStrategy.MINERU:
        return MinerUParser()

    # 默认按文件类型匹配
    for parser in sorted(PARSERS, key=lambda p: p.priority):
        if parser.can_parse(path):
            return parser
    raise ValueError(f"No parser found for {path}")


def parse_document(
    path: str | Path,
    hint: ParseStrategy | None = None,
    keywords: list[str] | None = None,
) -> Document:
    """主入口：解析单个文档（业务层 1 行调用）

    Args:
        path: 文件路径
        hint: 强制使用某档 PDF 策略（默认自动选择）
        keywords: 业务关键字列表（验收用）

    Returns:
        Document 对象（含 chunks + 验收结果在 metadata）
    """
    path = Path(path)
    if not path.exists():
        return Document(
            file_path=str(path),
            file_type=FileType.UNKNOWN,
            error=f"File not found: {path}",
        )

    parser = route_parser(path, hint)
    doc = parser.parse(path)

    # 业务验收
    if keywords and not doc.error:
        validation = doc.validate(keywords)
        doc.metadata["validation"] = validation

    return doc


# =============================================================================
# CLI 测试
# =============================================================================

if __name__ == "__main__":
    import sys
    if len(sys.argv) < 2:
        print("Usage: python code-doc-parse-pipeline.py <file_path> [keywords...]")
        sys.exit(1)

    file_path = Path(sys.argv[1])
    keywords = sys.argv[2:] if len(sys.argv) > 2 else None

    doc = parse_document(file_path, keywords=keywords)
    print(f"File: {doc.file_path}")
    print(f"Type: {doc.file_type}")
    print(f"Strategy: {doc.strategy}")
    print(f"Pages: {doc.page_count}")
    print(f"Chunks: {len(doc.chunks)}")
    print(f"Tables: {len(doc.table_chunks)}")
    print(f"Parse time: {doc.parse_time_ms}ms")
    if doc.error:
        print(f"Error: {doc.error}")
    if "validation" in doc.metadata:
        print(f"Validation: {doc.metadata['validation']}")
    print("---")
    print(doc.text[:500])
